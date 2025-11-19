"""
API - 语音服务网关 (完整版)
包含所有功能：转写、会议纪要、Dify集成、OpenAI兼容等
"""

import os
import uuid
import logging
import threading
import json
import asyncio
import zipfile
import io
import base64
from datetime import datetime
from typing import Optional, List, Union
from concurrent.futures import ThreadPoolExecutor, wait
import jieba
import jieba.analyse

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Request, WebSocket, WebSocketDisconnect
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openai import OpenAI

from application.voice.pipeline_service_funasr import PipelineService  # 使用FunASR版本
from infra.audio_io.storage import AudioStorage
from infra.websocket import ws_manager
from config import FILE_CONFIG, LANGUAGE_CONFIG

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/voice", tags=["voice"])

# 全局变量
pipeline_service: Optional[PipelineService] = None
audio_storage: Optional[AudioStorage] = None

# 历史记录文件
HISTORY_FILE = os.path.join(FILE_CONFIG['output_dir'], 'history_records.json')

# 线程安全的文件管理器
class ThreadSafeFileManager:
    """线程安全的文件管理器"""
    
    def __init__(self):
        self._files = []
        self._processing_files = []
        self._completed_files = []
        self._lock = threading.RLock()  # 递归锁，支持同一线程多次获取
    
    def add_file(self, file_info: dict):
        """添加文件"""
        with self._lock:
            self._files.append(file_info)
    
    def get_file(self, file_id: str) -> Optional[dict]:
        """获取文件信息"""
        with self._lock:
            for f in self._files:
                if f['id'] == file_id:
                    return f
            return None
    
    def get_all_files(self) -> List[dict]:
        """获取所有文件（返回副本）"""
        with self._lock:
            return self._files.copy()
    
    def update_file(self, file_id: str, updates: dict):
        """更新文件信息"""
        with self._lock:
            for f in self._files:
                if f['id'] == file_id:
                    f.update(updates)
                    return True
            return False
    
    def remove_file(self, file_id: str) -> bool:
        """移除文件"""
        with self._lock:
            for i, f in enumerate(self._files):
                if f['id'] == file_id:
                    self._files.pop(i)
                    self._processing_files = [fid for fid in self._processing_files if fid != file_id]
                    self._completed_files = [fid for fid in self._completed_files if fid != file_id]
                    return True
            return False
    
    def add_to_processing(self, file_id: str):
        """添加到处理队列"""
        with self._lock:
            if file_id not in self._processing_files:
                self._processing_files.append(file_id)
    
    def remove_from_processing(self, file_id: str):
        """从处理队列移除"""
        with self._lock:
            self._processing_files = [fid for fid in self._processing_files if fid != file_id]
    
    def add_to_completed(self, file_id: str):
        """添加到已完成队列"""
        with self._lock:
            if file_id not in self._completed_files:
                self._completed_files.append(file_id)
    
    def get_processing_files(self) -> List[str]:
        """获取处理中的文件ID列表"""
        with self._lock:
            return self._processing_files.copy()
    
    def get_completed_files(self) -> List[str]:
        """获取已完成的文件ID列表"""
        with self._lock:
            return self._completed_files.copy()
    
    def to_dict(self) -> dict:
        """转换为字典（用于序列化）"""
        with self._lock:
            return {
                'files': self._files.copy(),
                'processing_files': self._processing_files.copy(),
                'completed_files': self._completed_files.copy()
            }

# 使用线程安全的文件管理器
uploaded_files_manager = ThreadSafeFileManager()

# 线程池用于并发处理转写任务（从配置读取）
from config import CONCURRENCY_CONFIG
TRANSCRIPTION_THREAD_POOL = ThreadPoolExecutor(
    max_workers=CONCURRENCY_CONFIG.get('transcription_workers', 5),
    thread_name_prefix='transcribe-worker'
)

# 任务字典：存储 file_id -> Future 的映射，用于取消任务
transcription_tasks = {}  # {file_id: Future}
transcription_tasks_lock = threading.Lock()  # 保护任务字典的锁

# ⚠️ 移除全局锁 - 模型池已经处理并发，不再需要全局锁


# 保存主事件循环引用
_main_loop = None

def set_main_loop(loop):
    """设置主事件循环引用"""
    global _main_loop
    _main_loop = loop
    logger.info("主事件循环已设置")

def send_ws_message_sync(file_id: str, status: str, progress: int = 0, message: str = "", **kwargs):
    """
    在同步代码中发送WebSocket消息的辅助函数
    通过asyncio.run_coroutine_threadsafe在事件循环中执行异步任务
    """
    if _main_loop is None:
        logger.warning("主事件循环未设置，无法发送WebSocket消息")
        return
    
    try:
        # 在主事件循环中调度异步任务
        asyncio.run_coroutine_threadsafe(
            ws_manager.send_file_status(file_id, status, progress, message, kwargs),
            _main_loop
        )
    except Exception as e:
        logger.error(f"发送WebSocket消息失败: {e}")


def init_voice_gateway(service: PipelineService, storage: AudioStorage):
    """初始化网关服务"""
    global pipeline_service, audio_storage
    pipeline_service = service
    audio_storage = storage
    # 启动时加载历史记录
    load_history_from_file()


def load_history_from_file():
    """从文件加载历史记录（只加载已完成的，不影响当前正在处理的文件）"""
    global uploaded_files_manager
    try:
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                completed_files_from_disk = data.get('files', [])
                
                # 保留当前内存中未完成的文件
                all_files = uploaded_files_manager.get_all_files()
                current_incomplete_files = [f for f in all_files 
                                           if f['status'] in ['uploaded', 'processing', 'error']]
                
                # 合并：未完成的文件 + 磁盘上的已完成文件
                # 使用字典去重，以file_id为key
                files_dict = {}
                
                # 先添加未完成的文件
                for f in current_incomplete_files:
                    files_dict[f['id']] = f
                
                # 再添加已完成的文件（如果有重复，已完成的会覆盖）
                for f in completed_files_from_disk:
                    files_dict[f['id']] = f
                
                # 重新构建管理器（需要在锁内完成）
                uploaded_files_manager._lock.acquire()
                try:
                    uploaded_files_manager._files = list(files_dict.values())
                    uploaded_files_manager._completed_files = data.get('completed_files', [])
                finally:
                    uploaded_files_manager._lock.release()
                
                logger.info(f"已加载 {len(completed_files_from_disk)} 条历史记录，当前总文件数: {len(files_dict)}")
    except Exception as e:
        logger.error(f"加载历史记录失败: {e}")


def save_history_to_file():
    """保存历史记录到文件"""
    try:
        # 只保存已完成的文件记录
        all_files = uploaded_files_manager.get_all_files()
        completed_files = [f for f in all_files if f['status'] == 'completed']
        data = {
            'files': completed_files,
            'completed_files': uploaded_files_manager.get_completed_files()
        }
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info(f"已保存 {len(completed_files)} 条历史记录")
    except Exception as e:
        logger.error(f"保存历史记录失败: {e}")


def allowed_file(filename: str) -> bool:
    """检查文件格式"""
    ALLOWED_EXTENSIONS = {'mp3', 'wav', 'm4a', 'flac', 'aac', 'ogg', 'wma'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def save_transcript_to_word(transcript_data, filename_prefix="transcript", language="zh", audio_filename=None, file_id=None):
    """将转录结果保存为Word文档"""
    try:
        doc = Document()
        
        # 定义黑色（RGB(0,0,0)）
        black_color = RGBColor(0, 0, 0)
        
        title = doc.add_heading('语音转文字结果', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置标题为微软雅黑，黑色
        for run in title.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.color.rgb = black_color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_paragraph()
        
        info_table = doc.add_table(rows=3, cols=2)
        # 恢复原来的表格样式
        info_table.style = 'Light Grid Accent 1'
        
        for row in info_table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(5.0)
        
        # 设置表格第一列（标签）为宋体11号加粗，黑色，居中
        info_table.rows[0].cells[0].text = '生成时间'
        label_para = info_table.rows[0].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.runs[0]
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.name = 'SimSun'
        label_run.font.color.rgb = black_color
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置表格第二列（值）为宋体11号加粗，黑色，居中
        info_table.rows[0].cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        value_para = info_table.rows[0].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.runs[0]
        value_run.bold = True
        value_run.font.size = Pt(11)
        value_run.font.name = 'SimSun'
        value_run.font.color.rgb = black_color
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        info_table.rows[1].cells[0].text = '音频文件'
        label_para = info_table.rows[1].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.runs[0]
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.name = 'SimSun'
        label_run.font.color.rgb = black_color
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        info_table.rows[1].cells[1].text = audio_filename or "未知文件"
        value_para = info_table.rows[1].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.runs[0]
        value_run.bold = True
        value_run.font.size = Pt(11)
        value_run.font.name = 'SimSun'
        value_run.font.color.rgb = black_color
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        info_table.rows[2].cells[0].text = '文本长度'
        label_para = info_table.rows[2].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.runs[0]
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.name = 'SimSun'
        label_run.font.color.rgb = black_color
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        total_chars = sum(len(entry['text']) for entry in transcript_data)
        info_table.rows[2].cells[1].text = f"{total_chars} 字符"
        value_para = info_table.rows[2].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.runs[0]
        value_run.bold = True
        value_run.font.size = Pt(11)
        value_run.font.name = 'SimSun'
        value_run.font.color.rgb = black_color
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        
        for entry in transcript_data:
            speaker_para = doc.add_paragraph()
            speaker_run = speaker_para.add_run(entry['speaker'])
            speaker_run.bold = True
            speaker_run.font.size = Pt(12)
            speaker_run.font.name = 'SimSun'
            speaker_run.font.color.rgb = black_color
            speaker_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 设置发言人段落的下间距为0，使内容紧跟在后面
            speaker_para.paragraph_format.space_after = Pt(0)
            
            # 减小发言人和内容的间距，设置段落间距为0
            text_para = doc.add_paragraph()
            text_para.paragraph_format.space_before = Pt(0)
            text_para.paragraph_format.space_after = Pt(0)
            text_run = text_para.add_run(entry['text'])
            text_run.font.size = Pt(12)
            text_run.font.name = 'SimSun'
            text_run.font.color.rgb = black_color
            text_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 不同发言人之间的间距保持正常
            doc.add_paragraph()
        
        # ✅ 修复：使用微秒级时间戳 + file_id 确保文件名唯一性
        # 如果两个文件在同一秒内完成，使用微秒可以区分
        # 如果提供了 file_id，也加入文件名中，进一步确保唯一性
        now = datetime.now()
        timestamp = now.strftime('%Y%m%d_%H%M%S_%f')  # 包含微秒
        
        # 如果提供了 file_id，使用前8个字符作为唯一标识
        if file_id:
            file_id_short = file_id.replace('-', '')[:8]  # 移除连字符，取前8位
            filename = f"{filename_prefix}_{timestamp}_{file_id_short}.docx"
        else:
            filename = f"{filename_prefix}_{timestamp}.docx"
        
        filepath = os.path.join(FILE_CONFIG['output_dir'], filename)
        
        doc.save(filepath)
        return filename, filepath
        
    except Exception as e:
        logger.error(f"保存Word文档失败: {e}")
        return None, None


def generate_meeting_summary(transcript_data):
    """使用AI生成会议纪要"""
    try:
        if not transcript_data:
            return None
        
        transcript_text = ""
        for entry in transcript_data:
            speaker = entry.get('speaker', '未知发言人')
            text = entry.get('text', '')
            transcript_text += f"{speaker}: {text}\n\n"
        
        api_key = os.getenv('DEEPSEEK_API_KEY', os.getenv('OPENAI_API_KEY'))
        api_base = os.getenv('DEEPSEEK_API_BASE', os.getenv('OPENAI_API_BASE', 'https://api.deepseek.com'))
        
        if not api_key:
            logger.warning("未配置API KEY，使用默认模板")
            return generate_default_summary(transcript_data)
        
        client = OpenAI(api_key=api_key, base_url=api_base)
        
        prompt = f"""请根据以下会议转录内容，生成一份结构化的会议纪要。

会议转录内容：
{transcript_text}

请严格按照以下格式生成会议纪要：

【表格形式】
会议主题：[根据会议内容总结主题]
会议时间：[从转录中提取或推断时间]
会议地点：[从转录中提取或推断地点]
主持人：[从转录中识别主持人]
记录人：系统自动生成
参与人员：[从转录中识别所有参与者，用顿号分隔]

【正文部分】
一、会议议题及讨论内容
二、行动清单（待办事项）
三、其他说明"""
        
        response = client.chat.completions.create(
            model=os.getenv('DEEPSEEK_MODEL', os.getenv('OPENAI_MODEL', 'deepseek-chat')),
            messages=[
                {"role": "system", "content": "你是一个专业的会议纪要助手"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        summary = {
            'raw_text': response.choices[0].message.content,
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'model': os.getenv('DEEPSEEK_MODEL', 'deepseek-chat'),
            'status': 'success'
        }
        
        return summary
        
    except Exception as e:
        logger.error(f"生成会议纪要失败: {e}")
        return {
            'raw_text': f"生成会议纪要时发生错误: {str(e)}",
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'status': 'error',
            'error': str(e)
        }


def generate_default_summary(transcript_data):
    """生成默认会议纪要"""
    speaker_stats = {}
    total_words = 0
    
    for entry in transcript_data:
        speaker = entry.get('speaker', '未知发言人')
        text = entry.get('text', '')
        
        if speaker not in speaker_stats:
            speaker_stats[speaker] = {'count': 0, 'words': 0}
        
        speaker_stats[speaker]['count'] += 1
        speaker_stats[speaker]['words'] += len(text)
        total_words += len(text)
    
    summary_text = f"""## 会议概要
本次会议共有{len(speaker_stats)}位参与者，会议记录共{len(transcript_data)}段发言，总计约{total_words}字。

## 参与人员
"""
    
    for speaker, stats in speaker_stats.items():
        summary_text += f"- {speaker}: 发言{stats['count']}次\n"
    
    return {
        'raw_text': summary_text,
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'model': 'default_template',
        'status': 'success'
    }


def save_meeting_summary_to_word(transcript_data, summary_data, filename_prefix="meeting_summary"):
    """将会议纪要保存为Word文档"""
    try:
        doc = Document()
        
        # 定义黑色（RGB(0,0,0)）
        black_color = RGBColor(0, 0, 0)
        
        title = doc.add_heading('会议纪要', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置标题为黑色
        for run in title.runs:
            run.font.color.rgb = black_color
        
        # 添加纪要内容，所有文本设置为黑色
        for line in summary_data.get('raw_text', '').split('\n'):
            para = doc.add_paragraph(line)
            for run in para.runs:
                run.font.color.rgb = black_color
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{filename_prefix}_{timestamp}.docx"
        filepath = os.path.join(FILE_CONFIG['output_dir'], filename)
        
        doc.save(filepath)
        return filename, filepath
        
    except Exception as e:
        logger.error(f"保存会议纪要Word文档失败: {e}")
        return None, None


# ==================== API路由 ====================

# ==================== 方案1: 一站式转写接口 ====================

@router.post("/transcribe_all")
async def transcribe_all(
    audio_files: Union[UploadFile, List[UploadFile]] = File(...),  # 支持单个或多个文件
    language: str = Form("zh"),
    hotword: str = Form(""),
    generate_summary: bool = Form(False),
    return_type: str = Form("json")  # json/file/both
):
    """
    🎯 一站式音频转写接口（修复版）

    功能：上传单个/多个音频 + 转写 + 生成纪要 + 返回结果

    参数：
    - audio_files: 音频文件列表（必填，支持单个或多个）
    - language: 语言类型 (zh/en/zh-en/zh-dialect)，默认 zh
    - hotword: 热词，空格分隔，默认为空
    - generate_summary: 是否生成会议纪要，默认 False
    - return_type: 返回类型 (json/file/both)，默认 json
      - json: 返回JSON格式的结果和下载链接
      - file: 直接返回Word文档（单文件）或ZIP压缩包（多文件）
      - both: 返回JSON格式，同时在JSON中包含文件的base64编码 ⭐

    返回：
    - return_type=json: 返回JSON格式的转写结果
    - return_type=file: 直接返回Word文档或ZIP文件
    - return_type=both: 返回JSON（包含转写结果 + 文件base64编码）
    """
    try:
        # 标准化输入：将单个文件转换为列表
        if isinstance(audio_files, UploadFile):
            audio_files = [audio_files]
        
        # 验证输入
        if not audio_files:
            return JSONResponse({'success': False, 'message': '没有选择文件'}, status_code=400)
        
        # 验证所有文件格式
        for audio_file in audio_files:
            if not audio_file.filename:
                return JSONResponse({'success': False, 'message': '存在空文件名的文件'}, status_code=400)
            if not allowed_file(audio_file.filename):
                return JSONResponse({
                    'success': False, 
                    'message': f'文件 {audio_file.filename} 格式不支持，支持: mp3, wav, m4a, flac, aac, ogg, wma'
                }, status_code=400)
        
        logger.info(f"[一站式转写] 接收到 {len(audio_files)} 个文件，返回类型: {return_type}")
        
        # ================== 阶段1: 先保存所有文件到磁盘 ==================
        logger.info(f"[一站式转写] 阶段1: 保存 {len(audio_files)} 个文件...")
        files_to_process = []
        
        for idx, audio_file in enumerate(audio_files):
            try:
                # 保存文件
                filename = secure_filename(audio_file.filename)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                name, ext = os.path.splitext(filename)
                safe_filename = f"{name}_{timestamp}_{idx}{ext}"
                
                contents = await audio_file.read()
                file_size = len(contents)
                filepath = audio_storage.save_uploaded_file(contents, safe_filename)
                file_id = str(uuid.uuid4())
                
                file_info = {
                    'id': file_id,
                    'filename': safe_filename,
                    'original_name': audio_file.filename,
                    'filepath': filepath,
                    'size': file_size,
                    'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'status': 'processing',
                    'progress': 0,
                    'language': language,
                    'generate_summary': generate_summary,
                    'hotword': hotword
                }
                
                uploaded_files_manager.add_file(file_info)
                uploaded_files_manager.add_to_processing(file_id)
                files_to_process.append(file_info)
                
                # WebSocket通知
                send_ws_message_sync(file_id, 'processing', 0, f"已上传，等待转写: {audio_file.filename}")
                logger.info(f"[一站式转写] 已保存 {idx+1}/{len(audio_files)}: {audio_file.filename}")
                
            except Exception as e:
                logger.error(f"[一站式转写] 保存文件 {audio_file.filename} 失败: {e}")
        
        if not files_to_process:
            return JSONResponse({'success': False, 'message': '所有文件保存失败'}, status_code=500)
        
        logger.info(f"[一站式转写] 阶段1完成，成功保存 {len(files_to_process)} 个文件")
        
        # ================== 阶段2: 使用线程池并发处理所有文件 ==================
        import time as _time_module
        batch_start_time = _time_module.time()
        
        logger.info(f"[一站式转写] 阶段2: 并发转写 {len(files_to_process)} 个文件")
        logger.info(f"[一站式转写] 线程池配置: max_workers={CONCURRENCY_CONFIG['transcription_workers']}, 模型池={CONCURRENCY_CONFIG['asr_pool_size']}")
        
        # 共享结果容器（线程安全）
        all_results = []
        all_word_files = []
        results_lock = threading.Lock()
        
        # 并发度量指标
        concurrent_metrics = {
            'active_count': 0,
            'completed_count': 0,
            'failed_count': 0,
            'max_concurrent': 0,
            'start_times': {},
            'completion_times': {}
        }
        metrics_lock = threading.Lock()
        
        # 定义单文件处理函数
        def process_single_file_for_batch(file_info):
            file_id = file_info['id']
            original_name = file_info['original_name']
            file_start_time = _time_module.time()
            
            try:
                # 检查是否已被取消
                if file_info.get('_cancelled', False):
                    logger.info(f"[一站式转写-并发] 文件 {file_id} 已被取消，跳过处理")
                    file_info['status'] = 'uploaded'
                    file_info['progress'] = 0
                    return
                
                # 📊 记录开始 - 更新并发度量
                with metrics_lock:
                    concurrent_metrics['active_count'] += 1
                    concurrent_metrics['start_times'][file_id] = file_start_time
                    if concurrent_metrics['active_count'] > concurrent_metrics['max_concurrent']:
                        concurrent_metrics['max_concurrent'] = concurrent_metrics['active_count']
                
                logger.info(f"[一站式转写-并发] 开始处理: {original_name} (线程: {threading.current_thread().name}, 当前并发: {concurrent_metrics['active_count']})")
                
                # 创建进度回调
                def update_progress(step, progress, message="", transcript_entry=None):
                    # 检查是否已被取消
                    if file_info.get('_cancelled', False):
                        logger.info(f"[一站式转写-并发] 检测到文件 {file_id} 已被取消，停止处理")
                        raise InterruptedError("转写任务已被取消")
                    
                    file_info['progress'] = progress
                    send_ws_message_sync(file_id, 'processing', progress, message or f"处理中: {step}")
                
                # 执行转写
                pipeline_service.set_callback(update_progress)
                transcript, _, _ = pipeline_service.execute_transcription(
                    file_info['filepath'],
                    hotword=file_info['hotword'],
                    language=file_info['language'],
                    instance_id=file_id,
                    cancellation_flag=lambda: file_info.get('_cancelled', False)  # 传递取消检查函数
                )
                
                # 检查是否在转写过程中被取消
                if file_info.get('_cancelled', False):
                    logger.info(f"[一站式转写-并发] 文件 {file_id} 在转写过程中被取消")
                    file_info['status'] = 'uploaded'
                    file_info['progress'] = 0
                    file_info['error_message'] = '转写已停止'
                    send_ws_message_sync(file_id, 'uploaded', 0, '转写已停止')
                    with metrics_lock:
                        concurrent_metrics['active_count'] -= 1
                    return
                
                if not transcript:
                    logger.warning(f"[一站式转写-并发] 文件 {original_name} 转写失败")
                    file_info['status'] = 'error'
                    file_info['error_message'] = '转写失败'
                    send_ws_message_sync(file_id, 'error', 0, '转写失败')
                    
                    # 📊 记录失败
                    with metrics_lock:
                        concurrent_metrics['active_count'] -= 1
                        concurrent_metrics['failed_count'] += 1
                    
                    with results_lock:
                        all_results.append({
                            'success': False,
                            'filename': original_name,
                            'file_id': file_id,
                            'error': '转写失败'
                        })
                    return
                
                file_info['transcript_data'] = transcript
                
                # 保存转写文档
                # ✅ 修复：传入 file_id 确保每个文件生成唯一的转写文档文件名
                transcript_filename, transcript_filepath = save_transcript_to_word(
                    transcript,
                    language=file_info['language'],
                    audio_filename=original_name,
                    file_id=file_id
                )
                
                if transcript_filename:
                    file_info['transcript_file'] = transcript_filepath
                    with results_lock:
                        all_word_files.append(transcript_filepath)
                    logger.info(f"[一站式转写-并发] 转写文档已保存: {transcript_filename}")
                
                # 生成会议纪要（可选）
                summary = None
                if file_info['generate_summary']:
                    logger.info(f"[一站式转写-并发] 生成会议纪要: {original_name}")
                    summary = generate_meeting_summary(transcript)
                    if summary:
                        file_info['meeting_summary'] = summary
                        name, _ = os.path.splitext(file_info['filename'])
                        summary_filename, summary_filepath = save_meeting_summary_to_word(
                            transcript, summary, filename_prefix=f"summary_{name}"
                        )
                        if summary_filepath:
                            with results_lock:
                                all_word_files.append(summary_filepath)
                            logger.info(f"[一站式转写-并发] 会议纪要已保存: {summary_filename}")
                
                # 更新状态
                file_info['status'] = 'completed'
                file_info['progress'] = 100
                file_info['complete_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                
                if file_id not in uploaded_files_manager.get_completed_files():
                    uploaded_files_manager.add_to_completed(file_id)
                
                if file_id in uploaded_files_manager.get_processing_files():
                    uploaded_files_manager.remove_from_processing(file_id)
                
                save_history_to_file()
                send_ws_message_sync(file_id, 'completed', 100, '转写完成')
                
                # 统计信息
                speakers = set(t.get('speaker', '') for t in transcript if t.get('speaker'))
                total_duration = transcript[-1].get('end_time', 0) if transcript else 0
                
                # 构建单个文件结果
                file_result = {
                    'success': True,
                    'file_id': file_id,
                    'filename': original_name,
                    'file_info': {
                        'id': file_id,
                        'filename': original_name,
                        'upload_time': file_info['upload_time'],
                        'complete_time': file_info['complete_time'],
                        'size': file_info['size'],
                        'language': file_info['language']
                    },
                    'transcript': transcript,
                    'download_urls': {
                        'audio': f"/api/voice/audio/{file_id}?download=1",
                        'transcript': f"/api/voice/download_transcript/{file_id}"
                    },
                    'statistics': {
                        'speakers_count': len(speakers),
                        'segments_count': len(transcript),
                        'total_duration': round(total_duration, 2),
                        'total_characters': sum(len(t.get('text', '')) for t in transcript),
                        'speakers': list(speakers)
                    }
                }
                
                if summary:
                    file_result['summary'] = summary
                    file_result['download_urls']['summary'] = f"/api/voice/download_summary/{file_id}"
                
                with results_lock:
                    all_results.append(file_result)
                
                # 📊 记录成功完成
                file_duration = _time_module.time() - file_start_time
                with metrics_lock:
                    concurrent_metrics['active_count'] -= 1
                    concurrent_metrics['completed_count'] += 1
                    concurrent_metrics['completion_times'][file_id] = file_duration
                
                logger.info(f"[一站式转写-并发] ✅ 完成: {original_name} (耗时: {file_duration:.1f}秒, 剩余活跃: {concurrent_metrics['active_count']})")
                
            except InterruptedError as e:
                # 处理中断异常
                logger.info(f"[一站式转写-并发] 文件 {file_id} 转写被中断: {e}")
                file_info['status'] = 'uploaded'
                file_info['progress'] = 0
                file_info['error_message'] = '转写已停止'
                send_ws_message_sync(file_id, 'uploaded', 0, '转写已停止')
                with metrics_lock:
                    concurrent_metrics['active_count'] -= 1
            except Exception as e:
                file_id = file_info['id']
                logger.error(f"[一站式转写-并发] ❌ 处理文件 {file_info['original_name']} 失败: {e}")
                
                # 如果是因为取消导致的异常，不标记为错误
                if file_info.get('_cancelled', False):
                    file_info['status'] = 'uploaded'
                    file_info['progress'] = 0
                    file_info['error_message'] = '转写已停止'
                    send_ws_message_sync(file_id, 'uploaded', 0, '转写已停止')
                    with metrics_lock:
                        concurrent_metrics['active_count'] -= 1
                else:
                    import traceback
                    traceback.print_exc()
                    
                    # 📊 记录失败
                    with metrics_lock:
                        concurrent_metrics['active_count'] -= 1
                        concurrent_metrics['failed_count'] += 1
                    
                    with results_lock:
                        all_results.append({
                            'success': False,
                            'filename': file_info['original_name'],
                            'file_id': file_id,
                            'error': str(e)
                        })
            finally:
                # 从任务字典中移除
                file_id = file_info['id']
                with transcription_tasks_lock:
                    if file_id in transcription_tasks:
                        del transcription_tasks[file_id]
        
        # 使用线程池并发提交所有任务
        futures = []
        for file_info in files_to_process:
            file_id = file_info['id']
            # 初始化取消标志
            file_info['_cancelled'] = False
            
            future = TRANSCRIPTION_THREAD_POOL.submit(process_single_file_for_batch, file_info)
            futures.append(future)
            
            # 将Future存储到任务字典中，用于取消任务
            with transcription_tasks_lock:
                transcription_tasks[file_id] = future
        
        logger.info(f"[一站式转写] 📤 已提交 {len(futures)} 个任务到线程池并发处理")

        # 等待所有任务完成
        # ⚠️ 修复：使用 asyncio.to_thread 避免阻塞事件循环
        def wait_futures():
            wait(futures, timeout=3600)  # 最多等待1小时（减少超时时间）

        try:
            await asyncio.to_thread(wait_futures)
        except Exception as e:
            logger.warning(f"[一站式转写] 等待任务完成时出错: {e}，继续处理已完成的任务")
        
        # 计算总耗时和性能指标
        batch_duration = _time_module.time() - batch_start_time
        
        logger.info(f"[一站式转写] ==================== 批处理完成 ====================")
        logger.info(f"[一站式转写] 📊 性能统计:")
        logger.info(f"[一站式转写]   - 总文件数: {len(files_to_process)}")
        logger.info(f"[一站式转写]   - 成功: {concurrent_metrics['completed_count']}, 失败: {concurrent_metrics['failed_count']}")
        logger.info(f"[一站式转写]   - 最大并发数: {concurrent_metrics['max_concurrent']}")
        logger.info(f"[一站式转写]   - 总耗时: {batch_duration:.2f}秒")
        
        if concurrent_metrics['completion_times']:
            avg_time = sum(concurrent_metrics['completion_times'].values()) / len(concurrent_metrics['completion_times'])
            max_time = max(concurrent_metrics['completion_times'].values())
            min_time = min(concurrent_metrics['completion_times'].values())
            logger.info(f"[一站式转写]   - 单文件平均: {avg_time:.2f}秒, 最快: {min_time:.2f}秒, 最慢: {max_time:.2f}秒")
            logger.info(f"[一站式转写]   - 并发加速比: {(avg_time * len(files_to_process)) / batch_duration:.2f}x")
        
        logger.info(f"[一站式转写] ==================================================")
        
        # 统计
        success_count = sum(1 for r in all_results if r.get('success'))
        failed_count = len(all_results) - success_count
        
        logger.info(f"[一站式转写] 全部完成: 成功 {success_count}/{len(audio_files)}")
        
        # 根据 return_type 返回不同格式
        if return_type == "file":
            # 直接返回文件
            if not all_word_files:
                return JSONResponse({
                    'success': False,
                    'message': '所有文件转写失败',
                    'results': all_results
                }, status_code=500)
            
            # 单个文件：直接返回Word文档
            if len(all_word_files) == 1:
                return FileResponse(
                    path=all_word_files[0],
                    filename=os.path.basename(all_word_files[0]),
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            
            # 多个文件：打包成ZIP返回
            zip_filename = f"transcripts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
            zip_path = os.path.join(FILE_CONFIG['output_dir'], zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in all_word_files:
                    if os.path.exists(file_path):
                        arcname = os.path.basename(file_path)
                        zipf.write(file_path, arcname)
                        logger.info(f"[一站式转写] 添加到ZIP: {arcname}")
            
            logger.info(f"[一站式转写] ZIP创建成功: {zip_filename}")
            
            return FileResponse(
                path=zip_path,
                filename=zip_filename,
                media_type='application/zip',
                headers={
                    'X-Success-Count': str(success_count),
                    'X-Failed-Count': str(failed_count),
                    'X-Total-Files': str(len(all_word_files))
                }
            )
        
        elif return_type == "both":
            # both模式：返回JSON + 文件base64编码
            logger.info(f"[一站式转写] BOTH模式：准备编码 {len(all_word_files)} 个文件")
            
            files_data = []
            
            # 单个文件时，直接编码Word文档
            if len(all_word_files) == 1 and all_word_files[0]:
                try:
                    with open(all_word_files[0], 'rb') as f:
                        file_content = f.read()
                        file_base64 = base64.b64encode(file_content).decode('utf-8')
                        files_data.append({
                            'filename': os.path.basename(all_word_files[0]),
                            'content_base64': file_base64,
                            'size': len(file_content),
                            'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        })
                        logger.info(f"[一站式转写] 已编码文件: {os.path.basename(all_word_files[0])}")
                except Exception as e:
                    logger.error(f"[一站式转写] 编码文件失败: {e}")
            
            # 多个文件时，打包成ZIP再编码
            elif len(all_word_files) > 1:
                try:
                    zip_filename = f"transcripts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
                    zip_path = os.path.join(FILE_CONFIG['output_dir'], zip_filename)
                    
                    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for file_path in all_word_files:
                            if os.path.exists(file_path):
                                arcname = os.path.basename(file_path)
                                zipf.write(file_path, arcname)
                    
                    with open(zip_path, 'rb') as f:
                        zip_content = f.read()
                        zip_base64 = base64.b64encode(zip_content).decode('utf-8')
                        files_data.append({
                            'filename': zip_filename,
                            'content_base64': zip_base64,
                            'size': len(zip_content),
                            'mime_type': 'application/zip'
                        })
                        logger.info(f"[一站式转写] 已编码ZIP文件: {zip_filename}, 大小: {len(zip_content)} bytes")
                except Exception as e:
                    logger.error(f"[一站式转写] 创建ZIP失败: {e}")
            
            # 返回JSON响应（包含转写结果和文件）
            response_data = {
                'success': success_count > 0,
                'message': f'处理完成: 成功 {success_count}, 失败 {failed_count}',
                'total_files': len(audio_files),
                'success_count': success_count,
                'failed_count': failed_count,
                'results': all_results,
                'files': files_data  # ⭐ 包含文件的base64编码
            }
            
            # 添加下载链接（方便用户直接下载）
            if files_data:
                response_data['download_urls'] = []
                for file_info in files_data:
                    response_data['download_urls'].append({
                        'filename': file_info['filename'],
                        'url': f"/api/voice/download_file/{file_info['filename']}",
                        'size': file_info['size']
                    })
            
            logger.info(f"[一站式转写] 返回数据: files数量={len(files_data)}")
            return response_data
        
        else:
            # return_type == "json"，返回JSON响应（包含性能指标）
            return {
                'success': success_count > 0,
                'message': f'处理完成: 成功 {success_count}, 失败 {failed_count}',
                'total_files': len(audio_files),
                'success_count': success_count,
                'failed_count': failed_count,
                'results': all_results,
                'performance': {
                    'batch_duration': round(batch_duration, 2),
                    'max_concurrent': concurrent_metrics['max_concurrent'],
                    'avg_file_time': round(sum(concurrent_metrics['completion_times'].values()) / len(concurrent_metrics['completion_times']), 2) if concurrent_metrics['completion_times'] else 0,
                    'speedup_ratio': round((sum(concurrent_metrics['completion_times'].values())) / batch_duration, 2) if batch_duration > 0 and concurrent_metrics['completion_times'] else 1.0
                }
            }
        
    except Exception as e:
        logger.error(f"[一站式转写] 处理失败: {e}")
        import traceback
        traceback.print_exc()
        
        # 清理资源
        if 'file_info' in locals() and file_info:
            file_info['status'] = 'error'
            file_info['error_message'] = str(e)
            if file_id in uploaded_files_manager.get_processing_files():
                uploaded_files_manager.remove_from_processing(file_id)
        
        return JSONResponse({
            'success': False,
            'message': f'处理失败: {str(e)}'
        }, status_code=500)


# ==================== 方案2: RESTful文件资源接口 ====================

@router.get("/files")
async def list_all_files(
    filepath: Optional[str] = None,
    status: Optional[str] = None,
    limit: Optional[int] = None,
    offset: int = 0,
    include_history: bool = False,
    download: int = 0
):
    """
    📋 列出所有文件（RESTful风格，方案2优化）
    
    查询参数：
    - filepath: 可选，如果提供则直接返回该路径的音频文件（类似 /api/voice/files/{file_id}）
    - status: 过滤状态 (uploaded/processing/completed/error)
    - limit: 返回数量限制
    - offset: 分页偏移量
    - include_history: 是否包含历史记录，默认False
    - download: 当提供filepath时，是否下载（0=预览，1=下载）
    
    返回：文件列表及统计信息，或音频文件（当提供filepath时）
    """
    try:
        # 如果提供了filepath，直接返回音频文件
        if filepath:
            # 安全检查：防止路径遍历攻击
            # 规范化路径并确保在允许的目录内
            normalized_path = os.path.normpath(filepath)
            
            # 检查路径是否在uploads目录内
            upload_dir = os.path.abspath(FILE_CONFIG['upload_dir'])
            file_full_path = os.path.abspath(normalized_path)
            
            # 确保文件路径在uploads目录内
            if not file_full_path.startswith(upload_dir):
                raise HTTPException(status_code=403, detail="文件路径不在允许的目录内")
            
            # 检查文件是否存在
            if not os.path.exists(file_full_path):
                raise HTTPException(status_code=404, detail="音频文件不存在")
            
            # 检查是否为文件（不是目录）
            if not os.path.isfile(file_full_path):
                raise HTTPException(status_code=400, detail="指定路径不是文件")
            
            # 获取文件名（用于下载时的文件名）
            filename = os.path.basename(file_full_path)
            
            if download == 1:
                return FileResponse(
                    file_full_path,
                    media_type='application/octet-stream',
                    filename=filename
                )
            else:
                return FileResponse(
                    file_full_path,
                    media_type='audio/mpeg'
                )
        
        # 如果需要历史记录，从文件加载
        if include_history:
            load_history_from_file()
        
        # 获取所有文件
        all_files = uploaded_files_manager.get_all_files()
        
        # 根据状态过滤
        if status:
            filtered_files = [f for f in all_files if f['status'] == status]
        else:
            filtered_files = all_files
        
        # 排序：processing > uploaded > completed > error
        status_priority = {'processing': 0, 'uploaded': 1, 'completed': 2, 'error': 3}
        filtered_files.sort(key=lambda x: (
            status_priority.get(x['status'], 999),
            x.get('upload_time', '')
        ), reverse=True)
        
        # 分页
        total_count = len(filtered_files)
        if limit:
            filtered_files = filtered_files[offset:offset+limit]
        else:
            filtered_files = filtered_files[offset:]
        
        # 🔧 为每个文件添加可访问的下载URL
        for file_info in filtered_files:
            # 添加音频下载链接
            if 'download_urls' not in file_info:
                file_info['download_urls'] = {}
            file_info['download_urls']['audio'] = f"/api/voice/audio/{file_info['id']}?download=1"
            
            # 添加转写文档下载链接（如果存在）
            if file_info.get('transcript_file'):
                file_info['download_urls']['transcript'] = f"/api/voice/download_transcript/{file_info['id']}"
            
            # 添加会议纪要下载链接（如果存在）
            if file_info.get('meeting_summary'):
                file_info['download_urls']['summary'] = f"/api/voice/download_summary/{file_info['id']}"
        
        # 统计信息
        status_counts = {
            'uploaded': len([f for f in all_files if f['status'] == 'uploaded']),
            'processing': len([f for f in all_files if f['status'] == 'processing']),
            'completed': len([f for f in all_files if f['status'] == 'completed']),
            'error': len([f for f in all_files if f['status'] == 'error'])
        }
        
        return {
            'success': True,
            'files': filtered_files,
            'pagination': {
                'total': total_count,
                'limit': limit,
                'offset': offset,
                'returned': len(filtered_files)
            },
            'statistics': status_counts
        }
        
    except Exception as e:
        logger.error(f"列出文件失败: {e}")
        return JSONResponse({
            'success': False,
            'message': f'获取文件列表失败: {str(e)}'
        }, status_code=500)


@router.get("/files/{file_id}")
async def get_file_detail(
    file_id: str,
    include_transcript: bool = False,
    include_summary: bool = False
):
    """
    📄 获取文件详情（RESTful风格，方案2优化）
    
    路径参数：
    - file_id: 文件ID
    
    查询参数：
    - include_transcript: 是否包含转写结果，默认False
    - include_summary: 是否包含会议纪要，默认False
    
    返回：文件详细信息
    """
    try:
        file_info = next((f for f in uploaded_files_manager.get_all_files() if f['id'] == file_id), None)
        
        if not file_info:
            raise HTTPException(status_code=404, detail='文件不存在')
        
        # 构建基本响应
        result = {
            'success': True,
            'file': {
                'id': file_info['id'],
                'filename': file_info.get('original_name', file_info.get('filename')),
                'size': file_info.get('size', 0),
                'status': file_info['status'],
                'progress': file_info.get('progress', 0),
                'language': file_info.get('language', 'zh'),
                'upload_time': file_info.get('upload_time'),
                'complete_time': file_info.get('complete_time'),
                'error_message': file_info.get('error_message', '')
            }
        }
        
        # 添加下载链接
        result['file']['download_urls'] = {
            'audio': f"/api/voice/audio/{file_id}?download=1"
        }
        
        if file_info.get('transcript_file'):
            result['file']['download_urls']['transcript'] = f"/api/voice/download_transcript/{file_id}"
        
        if file_info.get('meeting_summary'):
            result['file']['download_urls']['summary'] = f"/api/voice/download_summary/{file_id}"
        
        # 可选：包含转写结果
        if include_transcript and file_info['status'] == 'completed':
            transcript_data = file_info.get('transcript_data', [])
            result['transcript'] = transcript_data
            
            # 添加统计信息
            if transcript_data:
                speakers = set(t.get('speaker', '') for t in transcript_data if t.get('speaker'))
                result['statistics'] = {
                    'speakers_count': len(speakers),
                    'segments_count': len(transcript_data),
                    'total_characters': sum(len(t.get('text', '')) for t in transcript_data),
                    'speakers': list(speakers)
                }
        
        # 可选：包含会议纪要
        if include_summary and file_info.get('meeting_summary'):
            result['summary'] = file_info['meeting_summary']
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取文件详情失败: {e}")
        raise HTTPException(status_code=500, detail=f'获取文件详情失败: {str(e)}')


@router.patch("/files/{file_id}")
async def update_file(file_id: str, request: Request):
    """
    🔄 更新文件（RESTful风格，方案2优化）
    
    路径参数：
    - file_id: 文件ID
    
    请求体：
    - action: 操作类型 (retranscribe/generate_summary)
    - language: 语言（重新转写时）
    - hotword: 热词（重新转写时）
    
    返回：更新后的文件信息
    """
    try:
        body = await request.json()
        action = body.get('action')
        
        file_info = next((f for f in uploaded_files_manager.get_all_files() if f['id'] == file_id), None)
        
        if not file_info:
            raise HTTPException(status_code=404, detail='文件不存在')
        
        if action == 'retranscribe':
            # 重新转写
            if file_info['status'] == 'processing':
                raise HTTPException(status_code=400, detail='文件正在处理中')
            
            language = body.get('language', file_info.get('language', 'zh'))
            hotword = body.get('hotword', '')
            
            # 重置状态
            file_info['status'] = 'processing'
            file_info['progress'] = 0
            file_info['language'] = language
            
            # 提交转写任务
            def retranscribe_task():
                try:
                    def update_progress(step, progress, message="", transcript_entry=None):
                        file_info['progress'] = progress
                        send_ws_message_sync(file_id, 'processing', progress, message)
                    
                    # ✅ 执行转写（不再需要全局锁）
                    pipeline_service.set_callback(update_progress)
                    transcript, _, _ = pipeline_service.execute_transcription(
                        file_info['filepath'],
                        hotword=hotword,
                        language=language,
                        instance_id=file_id
                    )
                    
                    if transcript:
                        file_info['transcript_data'] = transcript
                        # ✅ 修复：传入 file_id 确保每个文件生成唯一的转写文档文件名
                        filename, filepath = save_transcript_to_word(
                            transcript, language=language,
                            audio_filename=file_info['original_name'],
                            file_id=file_id
                        )
                        if filename:
                            file_info['transcript_file'] = filepath
                        
                        file_info['status'] = 'completed'
                        file_info['progress'] = 100
                        file_info['complete_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        save_history_to_file()
                        send_ws_message_sync(file_id, 'completed', 100, '重新转写完成')
                    else:
                        file_info['status'] = 'error'
                        file_info['error_message'] = '重新转写失败'
                        send_ws_message_sync(file_id, 'error', 0, '重新转写失败')
                        
                except Exception as e:
                    logger.error(f"重新转写失败: {e}")
                    file_info['status'] = 'error'
                    file_info['error_message'] = str(e)
                    send_ws_message_sync(file_id, 'error', 0, f"重新转写失败: {str(e)}")
            
            TRANSCRIPTION_THREAD_POOL.submit(retranscribe_task)
            
            return {
                'success': True,
                'message': '已开始重新转写',
                'file_id': file_id,
                'status': 'processing'
            }
        
        elif action == 'generate_summary':
            # 生成会议纪要
            if file_info['status'] != 'completed':
                raise HTTPException(status_code=400, detail='文件转写未完成')
            
            transcript_data = file_info.get('transcript_data', [])
            if not transcript_data:
                raise HTTPException(status_code=400, detail='没有转写结果')
            
            summary = generate_meeting_summary(transcript_data)
            if summary:
                file_info['meeting_summary'] = summary
                save_history_to_file()
                return {
                    'success': True,
                    'message': '会议纪要生成成功',
                    'summary': summary
                }
            else:
                raise HTTPException(status_code=500, detail='生成会议纪要失败')
        
        else:
            raise HTTPException(status_code=400, detail=f'不支持的操作: {action}')
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"更新文件失败: {e}")
        raise HTTPException(status_code=500, detail=f'更新文件失败: {str(e)}')


# ==================== 原有接口（保持向后兼容） ====================

@router.post("/upload")
async def upload_audio(audio_file: UploadFile = File(...)):
    """上传音频文件"""
    if not audio_file.filename:
        return JSONResponse({'success': False, 'message': '没有选择文件'})
    
    if not allowed_file(audio_file.filename):
        return JSONResponse({'success': False, 'message': '不支持的文件格式'})
    
    try:
        filename = secure_filename(audio_file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        name, ext = os.path.splitext(filename)
        safe_filename = f"{name}_{timestamp}{ext}"
        
        contents = await audio_file.read()
        file_size = len(contents)
        filepath = audio_storage.save_uploaded_file(contents, safe_filename)
        
        file_id = str(uuid.uuid4())
        
        file_info = {
            'id': file_id,
            'filename': safe_filename,
            'original_name': audio_file.filename,
            'filepath': filepath,
            'size': file_size,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'status': 'uploaded',
            'progress': 0,
            'error_message': ''
        }
        
        uploaded_files_manager.add_file(file_info)
        logger.info(f"文件上传成功: {audio_file.filename}, ID: {file_id}")
        
        return {
            'success': True,
            'message': '文件上传成功',
            'file': file_info
        }
    except Exception as e:
        logger.error(f"文件上传失败: {e}")
        return JSONResponse({'success': False, 'message': f'文件保存失败: {str(e)}'})


@router.post("/transcribe")
async def transcribe(request: Request):
    """开始转写（支持批量和并发处理；支持等待完成再返回）"""
    global TRANSCRIPTION_THREAD_POOL
    
    try:
        body = await request.json()
        
        # ✅ 兼容模式：同时支持 file_id (单个) 和 file_ids (数组)
        file_ids = body.get('file_ids', [])
        file_id = body.get('file_id', '')
        
        # 如果提供了单个 file_id，转换为数组
        if file_id and not file_ids:
            file_ids = [file_id]
        # 如果 file_ids 是字符串（某些情况下），也转换为数组
        elif isinstance(file_ids, str):
            file_ids = [file_ids]
        
        language = body.get('language', 'zh')
        hotword = body.get('hotword', '')
        # 新增：是否等待完成以及超时时间（秒）
        wait_until_complete = body.get('wait', True)
        timeout_seconds = int(body.get('timeout', 3600))  # 默认最多等待1小时
    except:
        return {'success': False, 'message': '请求参数错误'}
    
    if not file_ids:
        return {'success': False, 'message': '请选择要转写的文件（file_id 或 file_ids）'}
    
    # 检查所有文件是否存在且可处理
    files_to_process = []
    for file_id in file_ids:
        file_info = next((f for f in uploaded_files_manager.get_all_files() if f['id'] == file_id), None)
        if file_info:
            if file_info['status'] == 'processing':
                return {'success': False, 'message': f'文件 {file_info["original_name"]} 正在处理中'}
            files_to_process.append(file_info)
        else:
            return {'success': False, 'message': f'文件ID {file_id} 不存在'}
    
    if not files_to_process:
        return {'success': False, 'message': '没有可处理的文件'}
    
    # 🔧 提前更新所有文件状态为 processing，这样前端立即可以看到状态变化
    for file_info in files_to_process:
        file_info['status'] = 'processing'
        file_info['progress'] = 0
        file_info['language'] = language
        uploaded_files_manager.add_to_processing(file_info['id'])
        logger.info(f"文件 {file_info['original_name']} 状态已更新为 processing")
        
        # 🔔 WebSocket推送：开始转写
        send_ws_message_sync(
            file_info['id'], 
            'processing', 
            0, 
            f"开始转写: {file_info['original_name']}"
        )
    
    # 定义单文件处理函数
    def process_single_file(file_info):
        try:
            file_id = file_info['id']
            logger.info(f"[线程池] 开始处理文件: {file_info['original_name']}, 线程: {threading.current_thread().name}")
            
            # 检查是否已被取消
            if file_info.get('_cancelled', False):
                logger.info(f"[线程池] 文件 {file_id} 已被取消，跳过处理")
                file_info['status'] = 'uploaded'
                file_info['progress'] = 0
                return
            
            # 创建进度回调
            def update_file_progress(step, progress, message="", transcript_entry=None):
                # 检查是否已被取消
                if file_info.get('_cancelled', False):
                    logger.info(f"[线程池] 检测到文件 {file_id} 已被取消，停止处理")
                    raise InterruptedError("转写任务已被取消")
                
                file_info['progress'] = progress
                # 🔔 WebSocket推送：进度更新
                send_ws_message_sync(
                    file_id,
                    'processing',
                    progress,
                    message or f"处理中: {step}"
                )
            
            # ✅ 不再需要全局锁 - 模型池已经处理并发
            # 设置回调
            pipeline_service.set_callback(update_file_progress)
            
            # 再次检查是否已被取消
            if file_info.get('_cancelled', False):
                logger.info(f"[线程池] 文件 {file_id} 在开始转写前已被取消")
                file_info['status'] = 'uploaded'
                file_info['progress'] = 0
                return
            
            logger.info(f"[线程池] 开始转写: {file_info['original_name']}")
            transcript, _, _ = pipeline_service.execute_transcription(
                file_info['filepath'],
                hotword=hotword,
                language=language,
                instance_id=file_id,
                cancellation_flag=lambda: file_info.get('_cancelled', False)  # 传递取消检查函数
            )
            
            # 检查是否在转写过程中被取消
            if file_info.get('_cancelled', False):
                logger.info(f"[线程池] 文件 {file_id} 在转写过程中被取消")
                file_info['status'] = 'uploaded'
                file_info['progress'] = 0
                file_info['error_message'] = '转写已停止'
                send_ws_message_sync(
                    file_id,
                    'uploaded',
                    0,
                    '转写已停止'
                )
                return
            
            logger.info(f"[线程池] 转写完成: {file_info['original_name']}")
            
            # 保存转写结果
            if transcript:
                file_info['transcript_data'] = transcript
                logger.info(f"[线程池] 已保存 {len(transcript)} 条转写记录")
                
                # 自动生成Word文档
                # ✅ 修复：传入 file_id 确保每个文件生成唯一的转写文档文件名
                filename, filepath = save_transcript_to_word(
                    transcript,
                    language=language,
                    audio_filename=file_info['original_name'],
                    file_id=file_id
                )
                if filename:
                    file_info['transcript_file'] = filepath
                    logger.info(f"[线程池] 转写文档已保存: {filename}")
                
                # 更新状态为完成
                file_info['status'] = 'completed'
                file_info['progress'] = 100
                file_info['complete_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                
                # 添加到已完成列表
                if file_info['id'] not in uploaded_files_manager.get_completed_files():
                    uploaded_files_manager.add_to_completed(file_info['id'])
                
                # 保存历史记录
                save_history_to_file()
                
                # 🔔 WebSocket推送：转写完成
                send_ws_message_sync(
                    file_info['id'],
                    'completed',
                    100,
                    f"转写完成: {file_info['original_name']}"
                )
                
                logger.info(f"[线程池] 文件处理完成: {file_info['original_name']}")
            else:
                file_info['status'] = 'error'
                file_info['error_message'] = '转写失败'
                
                # 🔔 WebSocket推送：转写失败
                send_ws_message_sync(
                    file_info['id'],
                    'error',
                    0,
                    '转写失败'
                )
                
        except InterruptedError as e:
            # 处理中断异常
            file_id = file_info['id']
            logger.info(f"[线程池] 文件 {file_id} 转写被中断: {e}")
            file_info['status'] = 'uploaded'
            file_info['progress'] = 0
            file_info['error_message'] = '转写已停止'
            
            # 🔔 WebSocket推送：转写已停止
            send_ws_message_sync(
                file_id,
                'uploaded',
                0,
                '转写已停止'
            )
        except Exception as e:
            file_id = file_info['id']
            logger.error(f"[线程池] 处理文件失败 {file_info['original_name']}: {e}")
            
            # 如果是因为取消导致的异常，不标记为错误
            if file_info.get('_cancelled', False):
                file_info['status'] = 'uploaded'
                file_info['progress'] = 0
                file_info['error_message'] = '转写已停止'
                send_ws_message_sync(
                    file_id,
                    'uploaded',
                    0,
                    '转写已停止'
                )
            else:
                file_info['status'] = 'error'
                file_info['error_message'] = str(e)
                
                # 🔔 WebSocket推送：异常错误
                send_ws_message_sync(
                    file_id,
                    'error',
                    0,
                    f"处理失败: {str(e)}"
                )
            
            import traceback
            traceback.print_exc()
        finally:
            file_id = file_info['id']
            # 从处理列表中移除
            if file_id in uploaded_files_manager.get_processing_files():
                uploaded_files_manager.remove_from_processing(file_id)
            
            # 从任务字典中移除
            with transcription_tasks_lock:
                if file_id in transcription_tasks:
                    del transcription_tasks[file_id]
    
    # 使用线程池并发处理所有文件，并保留 future 以便可选等待
    futures = []
    for file_info in files_to_process:
        file_id = file_info['id']
        # 初始化取消标志
        file_info['_cancelled'] = False
        
        future = TRANSCRIPTION_THREAD_POOL.submit(process_single_file, file_info)
        futures.append((future, file_info))
        
        # 将Future存储到任务字典中，用于取消任务
        with transcription_tasks_lock:
            transcription_tasks[file_id] = future
    
    logger.info(f"已提交 {len(files_to_process)} 个文件到线程池处理")
    
    # 如果需要阻塞等待至完成，则轮询等待直到完成或超时
    if wait_until_complete:
        import time as _time
        deadline = _time.time() + timeout_seconds
        pending_ids = set(fi['id'] for _, fi in futures)
        failed_ids = set()
        completed_ids = set()
        
        # 轮询状态直到全部完成或超时
        while _time.time() < deadline and pending_ids:
            finished_now = []
            for _, fi in futures:
                fid = fi['id']
                if fid not in pending_ids:
                    continue
                status = fi.get('status')
                if status in ('completed', 'error'):
                    finished_now.append(fid)
                    if status == 'completed':
                        completed_ids.add(fid)
                    else:
                        failed_ids.add(fid)
            for fid in finished_now:
                pending_ids.discard(fid)
            if pending_ids:
                _time.sleep(0.5)
        
        if pending_ids:
            # 有未完成任务（超时）
            return {
                'success': False,
                'message': '部分任务未在超时时间内完成',
                'completed_file_ids': sorted(list(completed_ids)),
                'failed_file_ids': sorted(list(failed_ids)),
                'pending_file_ids': sorted(list(pending_ids))
            }
        else:
            # 全部完成
            return {
                'success': True,
                'message': f'转写完成 {len(completed_ids)} 个文件',
                'file_ids': sorted(list(completed_ids))
            }
    
    # 非阻塞兼容模式：立即返回“已开始转写”
    return {
        'success': True,
        'message': f'已开始转写 {len(files_to_process)} 个文件',
        'file_ids': [f['id'] for f in files_to_process],
        'count': len(files_to_process)
    }


@router.post("/stop/{file_id}")
async def stop_transcription(file_id: str):
    """
    ⏹️ 停止转写（向后兼容接口）
    
    实现真正的任务中断：取消Future并设置中断标志
    """
    file_info = next((f for f in uploaded_files_manager.get_all_files() if f['id'] == file_id), None)
    
    if not file_info:
        return {'success': False, 'message': '文件不存在'}
    
    if file_info['status'] != 'processing':
        return {'success': False, 'message': '文件未在转写中'}
    
    # 设置中断标志
    file_info['_cancelled'] = True
    logger.info(f"🛑 设置文件 {file_id} 的中断标志")
    
    # 尝试取消Future任务
    with transcription_tasks_lock:
        if file_id in transcription_tasks:
            future = transcription_tasks[file_id]
            cancelled = future.cancel()
            if cancelled:
                logger.info(f"✅ 成功取消文件 {file_id} 的Future任务")
            else:
                logger.warning(f"⚠️ 文件 {file_id} 的Future任务无法取消（可能已开始执行）")
            # 从任务字典中移除
            del transcription_tasks[file_id]
    
    # 更新文件状态
    file_info['status'] = 'uploaded'
    file_info['progress'] = 0
    file_info['error_message'] = '转写已停止'
    
    if file_id in uploaded_files_manager.get_processing_files():
        uploaded_files_manager.remove_from_processing(file_id)
    
    # 🔔 WebSocket推送：转写已停止
    send_ws_message_sync(
        file_id,
        'uploaded',
        0,
        '转写已停止'
    )
    
    logger.info(f"🛑 已停止文件 {file_id} 的转写任务")
    return {'success': True, 'message': '已停止转写'}


@router.get("/status/{file_id}")
async def get_status(file_id: str):
    """
    📊 获取转写状态（向后兼容接口）
    
    推荐使用新接口: GET /api/voice/files/{file_id}
    """
    for f in uploaded_files_manager.get_all_files():
        if f['id'] == file_id:
            return {
                'success': True,
                'status': f['status'],
                'progress': f['progress'],
                'error_message': f.get('error_message', '')
            }
    
    return {'success': False, 'message': '文件不存在'}


@router.get("/result/{file_id}")
async def get_result(file_id: str):
    """
    📄 获取转写结果（向后兼容接口）
    
    推荐使用新接口: GET /api/voice/files/{file_id}?include_transcript=true&include_summary=true
    """
    for f in uploaded_files_manager.get_all_files():
        if f['id'] == file_id:
            if f['status'] != 'completed':
                return {'success': False, 'message': '文件转写未完成'}
            
            return {
                'success': True,
                'file_info': {
                    'id': f['id'],
                    'original_name': f['original_name'],
                    'upload_time': f['upload_time']
                },
                'transcript': f.get('transcript_data', []),
                'summary': f.get('meeting_summary')
            }
    
    return {'success': False, 'message': '文件不存在'}


@router.get("/history")
async def list_history():
    """
    📜 获取历史记录（向后兼容接口）
    
    推荐使用新接口: GET /api/voice/files?status=completed&include_history=true
    """
    # 从文件加载历史记录
    load_history_from_file()
    
    history_records = []
    for f in uploaded_files_manager.get_all_files():
        if f['status'] == 'completed':
            transcript_data = f.get('transcript_data', [])
            speakers = set(t.get('speaker', '') for t in transcript_data if t.get('speaker'))
            
            details = f"{len(speakers)}位发言人, {len(transcript_data)}段对话"
            
            history_records.append({
                'file_id': f['id'],
                'filename': f['original_name'],
                'transcribe_time': f.get('complete_time', f.get('upload_time', '-')),
                'status': 'completed',
                'details': details
            })
    
    history_records.sort(key=lambda x: x['transcribe_time'], reverse=True)
    
    logger.info(f"返回 {len(history_records)} 条历史记录")
    
    return {
        'success': True,
        'records': history_records,
        'total': len(history_records)
    }


@router.delete("/files/{file_id}")
async def delete_file(file_id: str):
    """
    🗑️ 删除文件（RESTful标准接口）
    
    删除音频文件、转写结果和相关文档
    
    特殊操作：
    - file_id = "_clear_dify": 清空dify生成的.zip文件和对应上传的音频
    - file_id = "_clear_all": 清空所有历史记录，包括所有转写文件以及所有音频
    """
    # 特殊操作：清空dify生成文件
    if file_id == "_clear_dify":
        try:
            deleted_count = 0
            deleted_zip_count = 0
            deleted_audio_count = 0
            deleted_transcript_count = 0
            
            # 收集需要删除的文件ID集合
            files_to_delete = set()
            transcript_files_to_delete = set()
            
            # 获取output_dir目录下的所有transcripts_开头的.zip文件（一站式转写接口生成的）
            output_dir = FILE_CONFIG['output_dir']
            if os.path.exists(output_dir):
                for filename in os.listdir(output_dir):
                    # 只处理transcripts_开头的.zip文件（一站式转写接口生成的）
                    if filename.startswith('transcripts_') and filename.endswith('.zip'):
                        zip_path = os.path.join(output_dir, filename)
                        try:
                            # 读取ZIP文件内容，找到对应的转写文档
                            with zipfile.ZipFile(zip_path, 'r') as zipf:
                                zip_file_list = zipf.namelist()
                                logger.info(f"ZIP文件 {filename} 包含 {len(zip_file_list)} 个文件")
                                
                                # 提取ZIP文件中的转写文档文件名
                                for zip_entry in zip_file_list:
                                    if zip_entry.endswith('.docx'):
                                        # 转写文档文件名格式：transcript_YYYYMMDD_HHMMSS.docx
                                        transcript_filename = zip_entry
                                        transcript_files_to_delete.add(transcript_filename)
                                        logger.info(f"找到转写文档: {transcript_filename}")
                            
                            # 删除ZIP文件
                            os.remove(zip_path)
                            deleted_zip_count += 1
                            logger.info(f"已删除dify生成的ZIP文件: {filename}")
                        except Exception as e:
                            logger.error(f"处理ZIP文件失败 {filename}: {e}")
            
            # 通过转写文档找到对应的音频文件
            all_files = uploaded_files_manager.get_all_files()
            
            for file_info in all_files:
                # 检查转写文档是否在要删除的列表中
                transcript_file = file_info.get('transcript_file')
                if transcript_file:
                    transcript_basename = os.path.basename(transcript_file)
                    # 检查转写文档是否在ZIP文件中
                    if transcript_basename in transcript_files_to_delete:
                        files_to_delete.add(file_info['id'])
                        logger.info(f"找到对应的音频文件: {file_info.get('original_name', 'unknown')} (ID: {file_info['id']})")
            
            # 删除找到的文件
            for file_id_to_delete in files_to_delete:
                file_info = uploaded_files_manager.get_file(file_id_to_delete)
                if not file_info:
                    continue
                
                try:
                    # 删除音频文件
                    if 'filepath' in file_info and os.path.exists(file_info['filepath']):
                        os.remove(file_info['filepath'])
                        deleted_audio_count += 1
                        logger.info(f"已删除音频文件: {file_info['filepath']}")
                    
                    # 删除转写文档
                    if 'transcript_file' in file_info and os.path.exists(file_info['transcript_file']):
                        os.remove(file_info['transcript_file'])
                        deleted_transcript_count += 1
                        logger.info(f"已删除转写文档: {file_info['transcript_file']}")
                    
                    # 删除会议纪要文档
                    if 'meeting_summary_file' in file_info and os.path.exists(file_info['meeting_summary_file']):
                        os.remove(file_info['meeting_summary_file'])
                        logger.info(f"已删除会议纪要文档: {file_info['meeting_summary_file']}")
                    
                    # 从内存中删除
                    uploaded_files_manager.remove_file(file_id_to_delete)
                    deleted_count += 1
                except Exception as e:
                    logger.error(f"删除文件失败 {file_info.get('original_name', 'unknown')}: {e}")
            
            # 保存更新后的历史记录
            save_history_to_file()
            
            logger.info(f"清空dify生成文件完成: 删除 {deleted_zip_count} 个ZIP文件, {deleted_audio_count} 个音频文件, {deleted_transcript_count} 个转写文档, {deleted_count} 条历史记录")
            
            return {
                'success': True, 
                'message': f'清空dify生成文件成功',
                'deleted': {
                    'zip_files': deleted_zip_count,
                    'audio_files': deleted_audio_count,
                    'transcript_files': deleted_transcript_count,
                    'records': deleted_count
                }
            }
        except Exception as e:
            logger.error(f"清空dify生成文件失败: {e}")
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f'清空dify生成文件失败: {str(e)}')
    
    # 特殊操作：清空所有历史记录
    elif file_id == "_clear_all":
        try:
            deleted_count = 0
            deleted_audio_count = 0
            deleted_transcript_count = 0
            
            # 获取所有文件
            all_files = uploaded_files_manager.get_all_files()
            
            for file_info in all_files:
                # 跳过正在处理中的文件
                if file_info['status'] == 'processing':
                    continue
                
                try:
                    # 删除音频文件
                    if 'filepath' in file_info and os.path.exists(file_info['filepath']):
                        os.remove(file_info['filepath'])
                        deleted_audio_count += 1
                        logger.info(f"已删除音频文件: {file_info['filepath']}")
                    
                    # 删除转写文档
                    if 'transcript_file' in file_info and os.path.exists(file_info['transcript_file']):
                        os.remove(file_info['transcript_file'])
                        deleted_transcript_count += 1
                        logger.info(f"已删除转写文档: {file_info['transcript_file']}")
                    
                    # 删除会议纪要文档
                    if 'meeting_summary_file' in file_info and os.path.exists(file_info['meeting_summary_file']):
                        os.remove(file_info['meeting_summary_file'])
                        logger.info(f"已删除会议纪要文档: {file_info['meeting_summary_file']}")
                    
                    # 从内存中删除
                    uploaded_files_manager.remove_file(file_info['id'])
                    deleted_count += 1
                except Exception as e:
                    logger.error(f"删除文件失败 {file_info.get('original_name', 'unknown')}: {e}")
            
            # 清空output_dir目录下的所有文件（包括.zip和.docx）
            output_dir = FILE_CONFIG['output_dir']
            if os.path.exists(output_dir):
                for filename in os.listdir(output_dir):
                    # 跳过history_records.json文件
                    if filename == 'history_records.json':
                        continue
                    file_path = os.path.join(output_dir, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                            logger.info(f"已删除输出文件: {filename}")
                    except Exception as e:
                        logger.error(f"删除输出文件失败 {filename}: {e}")
            
            # 清空历史记录文件（保留文件但清空内容）
            try:
                with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
                    json.dump({'files': [], 'completed_files': []}, f, ensure_ascii=False, indent=2)
                logger.info("已清空历史记录文件")
            except Exception as e:
                logger.error(f"清空历史记录文件失败: {e}")
            
            logger.info(f"清空所有历史记录完成: 删除 {deleted_audio_count} 个音频文件, {deleted_transcript_count} 个转写文档, {deleted_count} 条历史记录")
            
            return {
                'success': True, 
                'message': f'清空所有历史记录成功',
                'deleted': {
                    'audio_files': deleted_audio_count,
                    'transcript_files': deleted_transcript_count,
                    'records': deleted_count
                }
            }
        except Exception as e:
            logger.error(f"清空所有历史记录失败: {e}")
            raise HTTPException(status_code=500, detail=f'清空所有历史记录失败: {str(e)}')
    
    # 正常删除单个文件
    file_info = uploaded_files_manager.get_file(file_id)
    
    if not file_info:
        raise HTTPException(status_code=404, detail='文件不存在')
    
    # ✅ 修复：如果文件正在处理中，但已设置取消标志（停止转写），允许删除
    if file_info['status'] == 'processing' and not file_info.get('_cancelled', False):
        raise HTTPException(status_code=400, detail='文件正在处理中，无法删除')
    
    try:
        # 删除音频文件
        if os.path.exists(file_info['filepath']):
            os.remove(file_info['filepath'])
            logger.info(f"已删除音频文件: {file_info['filepath']}")
        
        # 删除转写文档（如果存在）
        if 'transcript_file' in file_info and os.path.exists(file_info['transcript_file']):
            os.remove(file_info['transcript_file'])
            logger.info(f"已删除转写文档: {file_info['transcript_file']}")
        
        # 删除会议纪要文档（如果存在）
        if 'meeting_summary_file' in file_info and os.path.exists(file_info['meeting_summary_file']):
            os.remove(file_info['meeting_summary_file'])
            logger.info(f"已删除会议纪要文档: {file_info['meeting_summary_file']}")
        
        # 从内存中删除（使用线程安全方法）
        uploaded_files_manager.remove_file(file_id)
        
        # 保存更新后的历史记录到磁盘
        save_history_to_file()
        
        # 🔔 WebSocket推送：文件已删除
        send_ws_message_sync(
            file_id,
            'deleted',
            0,
            f"文件已删除: {file_info['original_name']}"
        )
        
        logger.info(f"文件删除成功: {file_info['original_name']}, ID: {file_id}")
        
        return {'success': True, 'message': '文件删除成功'}
        
    except Exception as e:
        logger.error(f"删除文件失败: {e}")
        raise HTTPException(status_code=500, detail=f'删除文件失败: {str(e)}')


@router.get("/audio/{file_id}")
async def get_audio(file_id: str, download: int = 0):
    """获取音频文件"""
    file_info = next((f for f in uploaded_files_manager.get_all_files() if f['id'] == file_id), None)
    
    if not file_info:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    if not os.path.exists(file_info['filepath']):
        raise HTTPException(status_code=404, detail="音频文件不存在")
    
    if download == 1:
        return FileResponse(
            file_info['filepath'],
            media_type='application/octet-stream',
            filename=file_info['original_name']
        )
    else:
        return FileResponse(
            file_info['filepath'],
            media_type='audio/mpeg'
        )


@router.get("/download_transcript/{file_id}")
async def download_transcript(file_id: str):
    """下载转写结果"""
    file_info = next((f for f in uploaded_files_manager.get_all_files() if f['id'] == file_id), None)
    
    if not file_info:
        raise HTTPException(status_code=404, detail='文件不存在')
    
    if file_info['status'] != 'completed':
        raise HTTPException(status_code=400, detail='文件转写未完成')
    
    if 'transcript_file' in file_info and file_info['transcript_file']:
        filepath = file_info['transcript_file']
        if os.path.exists(filepath):
            return FileResponse(
                path=filepath,
                filename=os.path.basename(filepath),
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    
    transcript_data = file_info.get('transcript_data', [])
    if not transcript_data:
        raise HTTPException(status_code=400, detail='没有转写结果')
    
    # ✅ 修复：传入 file_id 确保每个文件生成唯一的转写文档文件名
    filename, filepath = save_transcript_to_word(
        transcript_data,
        language=file_info.get('language', 'zh'),
        audio_filename=file_info.get('original_name'),
        file_id=file_id
    )
    
    if filename and os.path.exists(filepath):
        file_info['transcript_file'] = filepath
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        raise HTTPException(status_code=500, detail='生成Word文档失败')


@router.post("/generate_summary/{file_id}")
async def generate_summary_legacy(file_id: str):
    """
    📝 生成会议纪要（向后兼容接口）
    
    推荐使用新接口: PATCH /api/voice/files/{file_id} with action=generate_summary
    """
    file_info = next((f for f in uploaded_files_manager.get_all_files() if f['id'] == file_id), None)
    
    if not file_info:
        return {'success': False, 'message': '文件不存在'}
    
    if file_info['status'] != 'completed':
        return {'success': False, 'message': '文件转写未完成'}
    
    transcript_data = file_info.get('transcript_data', [])
    if not transcript_data:
        return {'success': False, 'message': '没有转写结果'}
    
    summary = generate_meeting_summary(transcript_data)
    
    if summary:
        file_info['meeting_summary'] = summary
        save_history_to_file()
        return {'success': True, 'summary': summary}
    else:
        return {'success': False, 'message': '生成会议纪要失败'}


@router.get("/download_summary/{file_id}")
async def download_summary(file_id: str):
    """下载会议纪要"""
    file_info = next((f for f in uploaded_files_manager.get_all_files() if f['id'] == file_id), None)
    
    if not file_info:
        raise HTTPException(status_code=404, detail='文件不存在')
    
    if not file_info.get('meeting_summary'):
        raise HTTPException(status_code=400, detail='请先生成会议纪要')
    
    transcript_data = file_info.get('transcript_data', [])
    summary = file_info['meeting_summary']
    
    filename, filepath = save_meeting_summary_to_word(transcript_data, summary)
    
    if filename and os.path.exists(filepath):
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        raise HTTPException(status_code=500, detail='保存Word文档失败')


@router.get("/languages")
async def get_languages():
    """获取支持的语言列表"""
    return {
        'success': True,
        'languages': [
            {'value': key, 'name': value['name'], 'description': value['description']}
            for key, value in LANGUAGE_CONFIG.items()
        ]
    }


@router.get("/transcript_files")
async def list_transcript_files():
    """列出所有转写文件"""
    try:
        files = audio_storage.list_output_files('.docx')
        for f in files:
            stat = os.stat(f['filepath'])
            f['modified'] = datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
            f['type'] = 'Word文档'
        
        files.sort(key=lambda x: x['modified'], reverse=True)
        return {'success': True, 'files': files}
    except Exception as e:
        return {'success': False, 'message': str(e)}


@router.get("/download_file/{filename}")
async def download_file(filename: str):
    """
    📥 下载输出文件（Word文档、ZIP压缩包等）
    
    路径参数：
    - filename: 文件名（例如：transcripts_20251101_203654.zip）
    
    用途：
    - 下载 /transcribe_all 接口生成的 ZIP 压缩包
    - 下载单独的 Word 转写文档
    
    返回：文件流
    """
    try:
        filepath = os.path.join(FILE_CONFIG['output_dir'], filename)
        if os.path.exists(filepath):
            # 根据文件扩展名确定 MIME 类型
            if filename.endswith('.zip'):
                media_type = 'application/zip'
            elif filename.endswith('.docx'):
                media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                media_type = 'application/octet-stream'
            
            return FileResponse(
                filepath,
                media_type=media_type,
                filename=filename
            )
        else:
            raise HTTPException(status_code=404, detail="文件不存在")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/delete_file/{filename}")
async def delete_output_file(filename: str):
    """删除输出文件"""
    try:
        filepath = os.path.join(FILE_CONFIG['output_dir'], filename)
        if os.path.exists(filepath):
            os.remove(filepath)
            return {'success': True, 'message': '文件删除成功'}
        else:
            return {'success': False, 'message': '文件不存在'}
    except Exception as e:
        return {'success': False, 'message': f'删除失败: {str(e)}'}


@router.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    """
    WebSocket端点 - 实时推送文件状态更新
    
    客户端可以通过此WebSocket连接接收：
    - 文件上传状态
    - 转写进度更新
    - 转写完成通知
    - 错误提示
    """
    await ws_manager.connect(websocket)
    
    try:
        # 发送连接成功消息
        await websocket.send_json({
            "type": "connected",
            "message": "WebSocket连接已建立"
        })
        
        # 保持连接并处理客户端消息
        while True:
            data = await websocket.receive_text()
            # 可以处理客户端发送的消息（如订阅特定文件）
            try:
                message = json.loads(data)
                if message.get('type') == 'subscribe':
                    file_id = message.get('file_id')
                    if file_id:
                        ws_manager.subscribe_file(websocket, file_id)
                        await websocket.send_json({
                            "type": "subscribed",
                            "file_id": file_id,
                            "message": f"已订阅文件 {file_id} 的状态更新"
                        })
            except json.JSONDecodeError:
                pass
            
    except WebSocketDisconnect:
        ws_manager.disconnect(websocket)
        logger.info("WebSocket客户端断开连接")
    except Exception as e:
        logger.error(f"WebSocket错误: {e}")
        ws_manager.disconnect(websocket)


# ==================== Dify 专用防卡死接口 ====================

@router.post("/dify_safe_transcribe")
async def dify_safe_transcribe(
    audio_file: UploadFile = File(...),
    language: str = Form("zh"),
    hotword: str = Form("")
):
    """
    🔧 Dify 安全转写接口 - 防卡死版本

    特点：
    1. 只处理单个文件
    2. 立即返回响应，不等待转写完成
    3. 后台异步处理，确保前端不会卡死
    4. 返回任务ID供后续查询

    使用流程：
    1. 调用此接口上传文件
    2. 立即获得 task_id
    3. 使用 /api/voice/files/{task_id} 查询状态
    4. 完成后使用 /api/voice/result/{task_id} 获取结果
    """
    try:
        # 快速验证
        if not audio_file.filename:
            return JSONResponse({'success': False, 'message': '没有选择文件'}, status_code=400)

        if not allowed_file(audio_file.filename):
            return JSONResponse({'success': False, 'message': f'文件 {audio_file.filename} 格式不支持'}, status_code=400)

        logger.info(f"[Dify安全转写] 接收文件: {audio_file.filename}")

        # 保存文件
        try:
            filename = secure_filename(audio_file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            name, ext = os.path.splitext(filename)
            safe_filename = f"{name}_{timestamp}{ext}"

            contents = await audio_file.read()
            file_size = len(contents)
            filepath = audio_storage.save_uploaded_file(contents, safe_filename)
            task_id = str(uuid.uuid4())

            # 创建文件记录
            file_info = {
                'id': task_id,
                'filename': safe_filename,
                'original_name': audio_file.filename,
                'filepath': filepath,
                'size': file_size,
                'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'status': 'uploaded',
                'progress': 0,
                'language': language,
                'hotword': hotword
            }

            uploaded_files_manager.add_file(file_info)
            logger.info(f"[Dify安全转写] 文件已保存: {audio_file.filename}, ID: {task_id}")

        except Exception as e:
            logger.error(f"[Dify安全转写] 保存文件失败: {e}")
            return JSONResponse({'success': False, 'message': f'文件保存失败: {str(e)}'}, status_code=500)

        # 🎯 关键修复：立即返回，不开始转写
        # 让 Dify 前端立即收到响应，然后使用单独的接口开始转写
        return {
            'success': True,
            'message': '文件上传成功，可以开始转写',
            'task_id': task_id,
            'filename': audio_file.filename,
            'size': file_size,
            'status': 'uploaded',
            'next_step': {
                'action': 'start_transcription',
                'api': '/api/voice/transcribe',
                'method': 'POST',
                'body': {
                    'file_id': task_id,
                    'language': language,
                    'hotword': hotword,
                    'wait': False  # 关键：不等待完成
                }
            },
            'status_query': f'/api/voice/files/{task_id}',
            'result_query': f'/api/voice/result/{task_id}',
            'note': '前端不会卡死，文件已准备好转写'
        }

    except Exception as e:
        logger.error(f"[Dify安全转写] 处理失败: {e}")
        return JSONResponse({'success': False, 'message': f'处理失败: {str(e)}'}, status_code=500)

